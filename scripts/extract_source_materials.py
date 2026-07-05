#!/usr/bin/env python3
"""Extract local source materials into a raw extraction summary for AllinCMS packaging."""

from __future__ import annotations

import argparse
import csv
from datetime import datetime, timezone
from html.parser import HTMLParser
import json
import os
from pathlib import Path
import re
import sys
from typing import Any

from build_source_inventory import file_hash, validate_inventory
from build_source_site_package import slugify


MAX_TEXT_CHARS_DEFAULT = 12000
MAX_TABLE_ROWS_DEFAULT = 40
SENSITIVE_PATTERNS = (
    re.compile(r"\b(?:cookie|bearer|next-action|next-router-state-tree)\b", re.IGNORECASE),
)
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_LIKE_RE = re.compile(r"\+?\d[\d\s().-]{6,}\d")
CONTACT_KEY_RE = re.compile(
    r"(?:email|e-mail|contact|phone|tel|telephone|whatsapp|notification|recipient|publiccontact)",
    re.IGNORECASE,
)
PRODUCT_NAME_HEADERS = {"name", "product", "product name", "model", "model name", "title"}
PRODUCT_DESCRIPTION_HEADERS = {"description", "summary", "details", "application", "applications", "features"}
PRODUCT_SLUG_HEADERS = {"slug", "url slug", "handle"}
PRODUCT_CATEGORY_HEADERS = {"category", "categories", "product category", "type", "family"}
PRODUCT_TAG_HEADERS = {"tag", "tags", "keywords"}
NON_SPEC_HEADERS = PRODUCT_NAME_HEADERS | PRODUCT_DESCRIPTION_HEADERS | PRODUCT_SLUG_HEADERS | PRODUCT_CATEGORY_HEADERS | PRODUCT_TAG_HEADERS
RESERVED_NAV_LABELS = {
    "/": "Home",
    "/products": "Products",
    "/posts": "Posts",
    "/news": "News",
    "/blog": "Blog",
    "/contact": "Contact",
    "/contact-us": "Contact Us",
    "/about": "About",
    "/about-us": "About Us",
}
ARTICLE_TAG_STOPWORDS = {
    "a",
    "an",
    "and",
    "for",
    "how",
    "how to",
    "the",
    "to",
}


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def load_json(path: Path, label: str) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        raise SystemExit(f"ERROR: {label} not found: {path}") from None
    except json.JSONDecodeError as exc:
        raise SystemExit(f"ERROR: invalid {label}: {exc}") from None
    if not isinstance(data, dict):
        raise SystemExit(f"ERROR: {label} root must be an object")
    return data


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def article_tags_from_title(title: str) -> list[str]:
    cleaned = clean_cell(title).strip(" .:-")
    if not cleaned:
        return []
    candidates = [cleaned]
    # Keep only phrase-level tags. Splitting article titles on "for" or "and"
    # creates unusable fragments such as "a Warehouse".
    candidates.extend(split_multi_value(cleaned.replace(" and ", ",")))
    tags: list[str] = []
    for candidate in candidates:
        label = clean_cell(candidate).strip(" .:-")
        lower = label.lower()
        if not label or lower in ARTICLE_TAG_STOPWORDS:
            continue
        if len(label) < 6 or len(label.split()) < 2:
            continue
        if lower.startswith(("a ", "an ", "the ")):
            continue
        if label not in tags:
            tags.append(label)
    return tags[:6]


def skill_root() -> Path:
    return Path(__file__).resolve().parents[1]


def ensure_output_outside_skill(path: Path) -> None:
    resolved = path.resolve()
    root = skill_root().resolve()
    if resolved == root or root in resolved.parents:
        raise SystemExit("ERROR: raw extraction output must be stored outside the skill package")


def compact_text(value: str, max_chars: int) -> tuple[str, bool]:
    normalized = re.sub(r"\s+", " ", value).strip()
    for pattern in SENSITIVE_PATTERNS:
        if pattern.search(normalized):
            normalized = pattern.sub("[REDACTED]", normalized)
    normalized = EMAIL_RE.sub("[REDACTED_EMAIL]", normalized)
    normalized = PHONE_LIKE_RE.sub("[REDACTED_PHONE]", normalized)
    if len(normalized) > max_chars:
        return normalized[:max_chars].rstrip(), True
    return normalized, False


def clean_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def first_sentence(text: str, max_chars: int = 180) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return ""
    match = re.search(r"(.+?[.!?])(?:\s|$)", normalized)
    sentence = match.group(1).strip() if match else normalized
    return sentence[:max_chars].rstrip()


def content_block(text: str, max_chars: int = 900) -> list[dict[str, Any]]:
    compact, _ = compact_text(text, max_chars)
    return [{"type": "paragraph", "text": compact}] if compact else []


def ensure_min_text(text: str, minimum: int, additions: list[str], max_chars: int = 1200) -> str:
    parts = [clean_cell(text)]
    for addition in additions:
        if len(" ".join(part for part in parts if part)) >= minimum:
            break
        cleaned = clean_cell(addition)
        if cleaned and cleaned not in parts:
            parts.append(cleaned)
    joined = " ".join(part for part in parts if part).strip()
    return joined[:max_chars].rstrip()


def split_multi_value(value: str) -> list[str]:
    values = [item.strip() for item in re.split(r"[,;/|]+", value) if item.strip()]
    return values


def table_products(extractions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    products: list[dict[str, Any]] = []
    seen_slugs: set[str] = set()
    for extraction in extractions:
        if extraction.get("status") != "extracted":
            continue
        source_ref = str(extraction.get("sourceRef") or "")
        for table in extraction.get("tables", []) if isinstance(extraction.get("tables"), list) else []:
            if not isinstance(table, dict):
                continue
            rows = table.get("rows")
            if not isinstance(rows, list) or len(rows) < 2:
                continue
            header = [clean_cell(cell).lower() for cell in rows[0] if clean_cell(cell)]
            if not header:
                continue
            name_index = next((index for index, label in enumerate(header) if label in PRODUCT_NAME_HEADERS), -1)
            desc_index = next((index for index, label in enumerate(header) if label in PRODUCT_DESCRIPTION_HEADERS), -1)
            slug_index = next((index for index, label in enumerate(header) if label in PRODUCT_SLUG_HEADERS), -1)
            if name_index < 0 or desc_index < 0:
                continue
            for row in rows[1:]:
                if not isinstance(row, list):
                    continue
                name = clean_cell(row[name_index] if name_index < len(row) else "")
                description = clean_cell(row[desc_index] if desc_index < len(row) else "")
                if len(name) < 3 or len(description) < 12:
                    continue
                slug_value = clean_cell(row[slug_index] if 0 <= slug_index < len(row) else "")
                slug = slugify(slug_value or name)
                if slug in seen_slugs:
                    continue
                categories: list[str] = []
                tags: list[str] = []
                specs: list[dict[str, str]] = []
                for index, label in enumerate(header):
                    value = clean_cell(row[index] if index < len(row) else "")
                    if not value:
                        continue
                    if index in {name_index, desc_index, slug_index} or label in PRODUCT_NAME_HEADERS | PRODUCT_DESCRIPTION_HEADERS | PRODUCT_SLUG_HEADERS:
                        continue
                    if label in PRODUCT_CATEGORY_HEADERS:
                        categories.extend(split_multi_value(value))
                    elif label in PRODUCT_TAG_HEADERS:
                        tags.extend(split_multi_value(value))
                    elif label not in NON_SPEC_HEADERS:
                        specs.append({"label": label, "value": value})
                spec_sentence = " ".join(f"{spec['label']}: {spec['value']}." for spec in specs[:8])
                category_sentence = f"Category: {', '.join(categories)}." if categories else ""
                tag_sentence = f"Tags: {', '.join(tags)}." if tags else ""
                body = ensure_min_text(description, 110, [spec_sentence, category_sentence, tag_sentence])
                products.append(
                    {
                        "name": name,
                        "slug": slug,
                        "description": description,
                        "content": content_block(body, 1000),
                        "specs": specs,
                        "categories": categories,
                        "tags": tags,
                        "sourceRefs": [source_ref] if source_ref else [],
                    }
                )
                seen_slugs.add(slug)
                if len(products) >= 12:
                    return products
    return products


def sentence_after(text: str, start: int) -> str:
    snippet = text[start:].strip()
    match = re.match(r"(.+?\.)(?:\s|$)", snippet)
    return match.group(1).strip() if match else snippet[:180].strip()


def naturalize_plan_sentence(title: str, body: str) -> str:
    cleaned = clean_cell(body).rstrip(".")
    if not cleaned:
        return ""
    lower_title = title.lower()
    verb_map = {
        "introduce": "introduces",
        "explain": "explains",
        "provide": "provides",
        "present": "presents",
        "show": "shows",
        "highlight": "highlights",
        "collect": "collects",
        "cover": "covers",
        "ask": "asks",
    }
    first, _, rest = cleaned.partition(" ")
    verb = verb_map.get(first.lower())
    if verb and rest:
        rest = re.sub(r"\band\s+ask\b", "and asks", rest, flags=re.IGNORECASE)
        subject = "The home page" if lower_title == "home" else f"The {lower_title} page"
        return f"{subject} {verb} {rest}."
    return cleaned + "."


def planned_pages(extractions: list[dict[str, Any]], site_name: str, combined: str, source_refs: list[str]) -> list[dict[str, Any]]:
    pages_by_path: dict[str, dict[str, Any]] = {}

    def add_page(title: str, body: str, refs: list[str], purpose: str = "content_page", replace: bool = True) -> None:
        clean_title = title.strip()
        if not clean_title:
            return
        path = "/" if clean_title.lower() == "home" else "/" + slugify(clean_title)
        if path in pages_by_path and not replace:
            return
        natural_body = naturalize_plan_sentence(clean_title, body) or body
        page_body = ensure_min_text(
            natural_body,
            130,
            [
                f"This {clean_title.lower()} page gives visitors a concise path to understand the product scope, project support, and inquiry options.",
                "The copy focuses on buyer needs, practical selection criteria, and a clear next step for discussing project requirements.",
            ],
        )
        pages_by_path[path] = {
            "title": clean_title,
            "path": path,
            "purpose": "homepage" if path == "/" else purpose,
            "sections": [{"heading": clean_title if path != "/" else site_name, "body": page_body}],
            "sourceRefs": refs or source_refs,
        }

    fallback_body = combined[:800] if combined else "Source-backed homepage copy requires review."
    add_page("Home", fallback_body, source_refs, "homepage")
    for extraction in extractions:
        if extraction.get("status") != "extracted":
            continue
        text = clean_cell(extraction.get("text"))
        if not text:
            continue
        refs = [str(extraction.get("sourceRef"))] if extraction.get("sourceRef") else source_refs
        for match in re.finditer(r"\b(Home|About Us|About|Contact Us|Contact|Services|Solutions|OEM|Cases)\s+page\s+should\s+", text, re.IGNORECASE):
            title_raw = match.group(1).strip()
            title = "About Us" if title_raw.lower() == "about" else "Contact" if title_raw.lower() == "contact us" else title_raw.title()
            body = sentence_after(text, match.end())
            add_page(title, body or f"{title} page copy requires review.", refs)
    return list(pages_by_path.values())


def planned_articles(extractions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    posts: list[dict[str, Any]] = []
    seen_slugs: set[str] = set()
    for extraction in extractions:
        if extraction.get("status") != "extracted":
            continue
        if is_structured_content_json(extraction):
            continue
        raw_text = str(extraction.get("rawText") or extraction.get("text") or "")
        text = clean_cell(raw_text)
        source_ref = str(extraction.get("sourceRef") or "")
        markdown_posts = markdown_heading_articles(raw_text, source_ref, seen_slugs)
        if markdown_posts:
            posts.extend(markdown_posts)
            continue
        if "Article" not in text:
            continue
        pattern = re.compile(
            r"\bArticle\s+\d+\s*:\s*(.+?)(?:\.|\n)\s*(?:Cover\s+(.+?)(?=\bArticle\s+\d+\s*:|$))?",
            re.IGNORECASE,
        )
        for match in pattern.finditer(text):
            title = clean_cell(match.group(1)).rstrip(".")
            detail = clean_cell(match.group(2) or "")
            if len(title) < 8:
                continue
            slug = slugify(title)
            if slug in seen_slugs:
                continue
            body = ensure_min_text(
                detail or title,
                150,
                [
                    f"This article helps buyers evaluate {title.lower()} for real project requirements.",
                    "It highlights practical selection criteria, specification tradeoffs, and questions that project buyers can use before shortlisting products.",
                ],
            )
            posts.append(
                {
                    "title": title,
                    "slug": slug,
                    "excerpt": first_sentence(body, 220) or title,
                    "content": content_block(body, 1200),
                    "categories": ["Buying Guides"],
                    "tags": article_tags_from_title(title),
                    "sourceRefs": [source_ref] if source_ref else [],
                }
            )
            seen_slugs.add(slug)
    return posts


def markdown_heading_articles(raw_text: str, source_ref: str, seen_slugs: set[str]) -> list[dict[str, Any]]:
    posts: list[dict[str, Any]] = []
    matches = list(re.finditer(r"(?m)^#{1,3}\s+(.+?)\s*$", raw_text))
    for index, match in enumerate(matches):
        title = clean_cell(match.group(1)).rstrip(".")
        if len(title) < 8 or title.lower() in {"article ideas", "articles", "blog", "posts", "content plan"}:
            continue
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(raw_text)
        detail = clean_cell(raw_text[start:end])
        if len(detail) < 40:
            continue
        slug = slugify(title)
        if slug in seen_slugs:
            continue
        body = ensure_min_text(
            detail,
            150,
            [
                f"This article helps buyers evaluate {title.lower()} for real project requirements.",
                "It highlights practical selection criteria, specification tradeoffs, and questions that project buyers can use before shortlisting products.",
            ],
        )
        posts.append(
            {
                "title": title,
                "slug": slug,
                "excerpt": first_sentence(body, 220) or title,
                "content": content_block(body, 1200),
                "categories": ["Buying Guides"],
                "tags": article_tags_from_title(title),
                "sourceRefs": [source_ref] if source_ref else [],
            }
        )
        seen_slugs.add(slug)
    return posts


def is_structured_content_json(extraction: dict[str, Any]) -> bool:
    if extraction.get("type") != "json":
        return False
    structured = extraction.get("structured")
    if not isinstance(structured, dict):
        return False
    content_keys = {
        "site",
        "pages",
        "products",
        "posts",
        "forms",
        "media",
        "siteInfo",
        "navigation",
        "taxonomyPlan",
        "taxonomy",
        "mediaPolicy",
        "contactFormPolicy",
    }
    return any(key in structured for key in content_keys)


def text_posts(extractions: list[dict[str, Any]], site_name: str) -> list[dict[str, Any]]:
    posts: list[dict[str, Any]] = []
    seen_slugs: set[str] = set()
    for extraction in extractions:
        if extraction.get("status") != "extracted":
            continue
        if extraction.get("type") == "spreadsheet" or is_structured_content_json(extraction):
            continue
        source_ref = str(extraction.get("sourceRef") or "")
        text = clean_cell(extraction.get("text"))
        if len(text) < 240:
            continue
        if re.search(r"\bArticle\s+\d+\s*:", text, re.IGNORECASE):
            continue
        title = ""
        path = Path(str(extraction.get("path", "")))
        if path.stem:
            title = re.sub(r"[-_]+", " ", path.stem).strip().title()
        if not title or title.lower() in {"brief", "catalog", "source", "data"}:
            title = f"{site_name} Source Overview"
        slug = slugify(title)
        if slug in seen_slugs:
            continue
        excerpt = first_sentence(text, 220)
        posts.append(
            {
                "title": title,
                "slug": slug,
                "excerpt": excerpt,
                "content": content_block(text, 1400),
                "sourceRefs": [source_ref] if source_ref else [],
            }
        )
        seen_slugs.add(slug)
        if len(posts) >= 6:
            return posts
    return posts


class SimpleHTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return " ".join(self.parts)


def extract_text_file(path: Path, max_chars: int) -> dict[str, Any]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    text, truncated = compact_text(raw, max_chars)
    return {"text": text, "rawText": raw[:max_chars], "truncated": truncated, "tables": [], "method": "plain_text"}


def extract_html_file(path: Path, max_chars: int) -> dict[str, Any]:
    parser = SimpleHTMLTextExtractor()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    text, truncated = compact_text(parser.text(), max_chars)
    return {"text": text, "truncated": truncated, "tables": [], "method": "html_parser"}


def extract_json_file(path: Path, max_chars: int) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    text, truncated = compact_text(json.dumps(data, ensure_ascii=False), max_chars)
    structured = redact_sensitive_contact_fields(data) if isinstance(data, dict) else {}
    return {"text": text, "truncated": truncated, "tables": [], "method": "json_dump", "structured": structured}


def extract_pdf_file(path: Path, max_chars: int) -> dict[str, Any]:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        pdfplumber = None  # type: ignore
    if pdfplumber is not None:
        parts: list[str] = []
        page_count = 0
        with pdfplumber.open(str(path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        text, truncated = compact_text("\n".join(parts), max_chars)
        return {"text": text, "truncated": truncated, "tables": [], "method": "pdfplumber", "pageCount": page_count}
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    text, truncated = compact_text("\n".join(page.extract_text() or "" for page in reader.pages), max_chars)
    return {"text": text, "truncated": truncated, "tables": [], "method": "pypdf", "pageCount": len(reader.pages)}


def extract_docx_file(path: Path, max_chars: int, max_rows: int) -> dict[str, Any]:
    import docx  # type: ignore

    document = docx.Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows: list[list[str]] = []
        for row in table.rows[:max_rows]:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"tableIndex": table_index, "rows": rows, "truncated": len(table.rows) > max_rows})
    text, truncated = compact_text("\n".join(paragraphs), max_chars)
    return {"text": text, "truncated": truncated, "tables": tables, "method": "python-docx"}


def extract_xlsx_file(path: Path, max_chars: int, max_rows: int) -> dict[str, Any]:
    import openpyxl  # type: ignore

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    tables: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
            if row_index >= max_rows:
                break
            values = ["" if cell is None else str(cell) for cell in row]
            if any(value.strip() for value in values):
                rows.append(values)
                parts.append(" | ".join(values))
        tables.append({"sheet": sheet.title, "rows": rows, "truncated": sheet.max_row > max_rows})
    text, truncated = compact_text("\n".join(parts), max_chars)
    return {"text": text, "truncated": truncated, "tables": tables, "method": "openpyxl"}


def extract_csv_file(path: Path, max_chars: int, max_rows: int, delimiter: str = ",") -> dict[str, Any]:
    rows: list[list[str]] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for index, row in enumerate(reader):
            if index >= max_rows:
                break
            rows.append(row)
    text, truncated = compact_text("\n".join(" | ".join(row) for row in rows), max_chars)
    return {"text": text, "truncated": truncated, "tables": [{"rows": rows, "truncated": False}], "method": "csv"}


def extract_image_file(path: Path) -> dict[str, Any]:
    return {
        "text": "",
        "truncated": False,
        "tables": [],
        "method": "image_metadata_only",
        "mediaCandidate": {
            "path": str(path),
            "name": path.name,
            "requiresUserApprovalBeforeUpload": True,
        },
    }


def extract_entry(entry: dict[str, Any], args: argparse.Namespace) -> dict[str, Any]:
    path = Path(str(entry.get("path", "")))
    source_type = entry.get("type")
    base = {
        "sourceRef": entry.get("sourceRef"),
        "path": str(path) if str(path) else "",
        "name": entry.get("name"),
        "type": source_type,
        "sizeBytes": entry.get("sizeBytes"),
        "sha256": entry.get("sha256"),
    }
    if not path.exists():
        return {**base, "status": "failed", "hashVerified": False, "error": "file_not_found"}
    expected_sha = entry.get("sha256")
    current_sha = file_hash(path)
    if isinstance(expected_sha, str) and expected_sha and current_sha != expected_sha:
        return {
            **base,
            "status": "failed",
            "hashVerified": False,
            "currentSha256": current_sha,
            "error": "source_hash_mismatch",
        }
    current_size = path.stat().st_size
    expected_size = entry.get("sizeBytes")
    if isinstance(expected_size, int) and expected_size != current_size:
        return {
            **base,
            "status": "failed",
            "hashVerified": False,
            "currentSizeBytes": current_size,
            "error": "source_size_mismatch",
        }
    try:
        if source_type in {"text", "markdown"}:
            extracted = extract_text_file(path, args.max_text_chars)
        elif source_type == "html":
            extracted = extract_html_file(path, args.max_text_chars)
        elif source_type == "json":
            extracted = extract_json_file(path, args.max_text_chars)
        elif source_type == "pdf":
            extracted = extract_pdf_file(path, args.max_text_chars)
        elif source_type == "docx":
            extracted = extract_docx_file(path, args.max_text_chars, args.max_table_rows)
        elif source_type == "spreadsheet" and path.suffix.lower() in {".xlsx", ".xls"}:
            extracted = extract_xlsx_file(path, args.max_text_chars, args.max_table_rows)
        elif source_type == "spreadsheet" and path.suffix.lower() == ".tsv":
            extracted = extract_csv_file(path, args.max_text_chars, args.max_table_rows, delimiter="\t")
        elif source_type == "spreadsheet":
            extracted = extract_csv_file(path, args.max_text_chars, args.max_table_rows)
        elif source_type == "image":
            extracted = extract_image_file(path)
        else:
            return {**base, "status": "unsupported", "hashVerified": False, "error": f"unsupported type {source_type}"}
    except Exception as exc:
        return {**base, "status": "failed", "hashVerified": True, "error": f"{type(exc).__name__}: {exc}"}
    return {
        **base,
        "path": str(path),
        "sizeBytes": current_size,
        "sha256": current_sha,
        "hashVerified": True,
        "status": "extracted",
        **extracted,
    }


def bind_source_refs(value: Any, source_ref: str) -> Any:
    if isinstance(value, list):
        return [bind_source_refs(item, source_ref) for item in value]
    if not isinstance(value, dict):
        return value
    result: dict[str, Any] = {}
    for key, item in value.items():
        if key == "sourceRefs":
            result[key] = [source_ref] if source_ref else []
        else:
            result[key] = bind_source_refs(item, source_ref)
    if source_ref and any(key in result for key in ("title", "name", "sections", "description", "content")) and "sourceRefs" not in result:
        result["sourceRefs"] = [source_ref]
    return result


def redact_contact_value(value: str) -> dict[str, Any] | str:
    cleaned = clean_cell(value)
    if not cleaned:
        return ""
    if EMAIL_RE.search(cleaned):
        return {
            "status": "provided_in_source_redacted",
            "contactType": "email",
            "valueRedacted": True,
            "requiresUserConfirmation": True,
        }
    if PHONE_LIKE_RE.search(cleaned):
        return {
            "status": "provided_in_source_redacted",
            "contactType": "phone_or_messaging",
            "valueRedacted": True,
            "requiresUserConfirmation": True,
        }
    return cleaned


def redact_sensitive_contact_fields(value: Any, key_path: str = "") -> Any:
    if isinstance(value, list):
        return [redact_sensitive_contact_fields(item, key_path) for item in value]
    if isinstance(value, dict):
        redacted: dict[str, Any] = {}
        for key, item in value.items():
            next_path = f"{key_path}.{key}" if key_path else str(key)
            redacted[key] = redact_sensitive_contact_fields(item, next_path)
        return redacted
    if isinstance(value, str) and (CONTACT_KEY_RE.search(key_path) or EMAIL_RE.search(value) or PHONE_LIKE_RE.search(value)):
        return redact_contact_value(value)
    return value


def structured_objects(extractions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    objects: list[dict[str, Any]] = []
    for extraction in extractions:
        if extraction.get("status") != "extracted" or extraction.get("type") != "json":
            continue
        structured = extraction.get("structured")
        source_ref = str(extraction.get("sourceRef") or "")
        if isinstance(structured, dict):
            objects.append(normalize_structured_object(structured, source_ref))
    return objects


def object_nonempty(value: Any) -> bool:
    return value not in (None, "", [], {})


def normalize_structured_object(obj: dict[str, Any], default_source_ref: str = "") -> dict[str, Any]:
    normalized = dict(obj)
    site = normalized.get("site") if isinstance(normalized.get("site"), dict) else {}
    top_level_site = {
        key: normalized.get(key)
        for key in ("siteName", "siteDescription", "description", "language", "industry")
        if object_nonempty(normalized.get(key))
    }
    if top_level_site:
        normalized["site"] = {**site, **top_level_site}
    if "siteDescription" not in normalized.get("site", {}) and object_nonempty(normalized.get("description")):
        normalized.setdefault("site", {})["siteDescription"] = normalized["description"]

    if isinstance(normalized.get("requiredPages"), list) and not isinstance(normalized.get("pages"), list):
        normalized["pages"] = normalized["requiredPages"]

    content_goals = normalized.get("contentGoals")
    if isinstance(content_goals, dict):
        normalized["contentGoals"] = content_goals
    launch_deferrals = normalized.get("launchDeferrals")
    if isinstance(launch_deferrals, list):
        normalized.setdefault("openQuestions", [])
        if isinstance(normalized["openQuestions"], list):
            for item in launch_deferrals:
                value = clean_cell(item)
                if value:
                    normalized["openQuestions"].append(f"Launch deferral requires user confirmation: {value}")

    source_ref = default_source_ref
    if source_ref:
        normalized = bind_source_refs(normalized, source_ref)
    return redact_sensitive_contact_fields(normalized)


def merge_dicts(*values: Any) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for value in values:
        if isinstance(value, dict):
            merged.update(value)
    return merged


def merge_list_fields(objects: list[dict[str, Any]], key: str) -> list[Any]:
    merged: list[Any] = []
    seen: set[str] = set()
    for obj in objects:
        value = obj.get(key)
        if not isinstance(value, list):
            continue
        for item in value:
            marker = json.dumps(item, ensure_ascii=False, sort_keys=True)
            if marker in seen:
                continue
            merged.append(item)
            seen.add(marker)
    return merged


def nav_label_for_path(path: str) -> str:
    normalized = path.rstrip("/") or "/"
    if normalized in RESERVED_NAV_LABELS:
        return RESERVED_NAV_LABELS[normalized]
    return normalized.strip("/").replace("-", " ").replace("_", " ").title() or "Home"


def normalize_navigation_item(item: Any) -> dict[str, str] | None:
    if isinstance(item, str):
        raw = clean_cell(item)
        if not raw:
            return None
        path = raw if raw.startswith("/") else "/" + slugify(raw)
        return {"label": nav_label_for_path(path), "path": path}
    if not isinstance(item, dict):
        return None
    raw_path = clean_cell(item.get("path") or item.get("href") or item.get("url"))
    label = clean_cell(item.get("label") or item.get("title") or item.get("name"))
    if not raw_path and label:
        raw_path = "/" if label.lower() == "home" else "/" + slugify(label)
    if not raw_path:
        return None
    path = raw_path if raw_path.startswith("/") else "/" + raw_path
    return {"label": label or nav_label_for_path(path), "path": path}


def normalize_page_item(item: Any, source_refs: list[str]) -> dict[str, Any] | None:
    if isinstance(item, str):
        title = clean_cell(item)
        body = f"{title} page gives visitors a source-backed overview and keeps final details pending user confirmation."
        raw_path = "/" if title.lower() == "home" else "/" + slugify(title)
        return {
            "title": title or "Page",
            "path": raw_path,
            "purpose": "homepage" if raw_path == "/" else "content_page",
            "sections": [{"heading": title or "Page", "body": body}],
            "sourceRefs": source_refs,
        }
    if not isinstance(item, dict):
        return None
    title = clean_cell(item.get("title") or item.get("name") or item.get("label"))
    path = clean_cell(item.get("path") or item.get("route"))
    if not title and path:
        title = nav_label_for_path(path)
    if not title:
        return None
    if not path:
        path = "/" if title.lower() == "home" else "/" + slugify(title)
    if not path.startswith("/"):
        path = "/" + path
    body = clean_cell(item.get("body") or item.get("description") or item.get("purpose"))
    body = ensure_min_text(
        body or f"{title} page supports the source-backed site plan.",
        130,
        [
            "The page should summarize buyer-relevant information, connect to the product catalog or article library, and keep unconfirmed contact, media, pricing, and legal details deferred.",
        ],
    )
    page = dict(item)
    page.update(
        {
            "title": title,
            "path": path,
            "purpose": clean_cell(item.get("purpose")) or ("homepage" if path == "/" else "content_page"),
            "sections": item.get("sections") if isinstance(item.get("sections"), list) else [{"heading": title, "body": body}],
            "sourceRefs": item.get("sourceRefs") if isinstance(item.get("sourceRefs"), list) else source_refs,
        }
    )
    return page


def merge_site(objects: list[dict[str, Any]], fallback: dict[str, Any]) -> dict[str, Any]:
    merged = dict(fallback)
    for obj in objects:
        site = obj.get("site")
        if isinstance(site, dict):
            merged.update({key: value for key, value in site.items() if value not in (None, "", [], {})})
    return merged


def merge_navigation(objects: list[dict[str, Any]]) -> dict[str, Any]:
    items: list[Any] = []
    seen_paths: set[str] = set()
    for obj in objects:
        navigation = obj.get("navigation")
        if isinstance(navigation, dict):
            nav_items = navigation.get("items")
        elif isinstance(navigation, list):
            nav_items = navigation
        else:
            nav_items = None
        if not isinstance(nav_items, list):
            continue
        for item in nav_items:
            normalized = normalize_navigation_item(item)
            if not normalized:
                continue
            path = clean_cell(normalized.get("path"))
            label = clean_cell(normalized.get("label"))
            marker = path or label
            if not marker or marker in seen_paths:
                continue
            items.append(normalized)
            seen_paths.add(marker)
    return {"items": items} if items else {}


def merge_pages(objects: list[dict[str, Any]]) -> list[Any]:
    pages: list[Any] = []
    seen_paths: set[str] = set()
    for obj in objects:
        source_refs = []
        for ref in obj.get("sourceRefs", []) if isinstance(obj.get("sourceRefs"), list) else []:
            if isinstance(ref, str) and ref.strip():
                source_refs.append(ref)
        for item in obj.get("pages", []) if isinstance(obj.get("pages"), list) else []:
            page = normalize_page_item(item, source_refs)
            if not page:
                continue
            path = clean_cell(page.get("path"))
            if path in seen_paths:
                continue
            pages.append(page)
            seen_paths.add(path)
    return pages


def merge_taxonomy_plan(objects: list[dict[str, Any]]) -> dict[str, Any]:
    merged: dict[str, list[Any]] = {}
    for key in ("productCategories", "postCategories", "productTags", "postTags"):
        values: list[Any] = []
        seen: set[str] = set()
        for obj in objects:
            plan = obj.get("taxonomyPlan") or obj.get("taxonomy")
            if not isinstance(plan, dict):
                continue
            raw_values = plan.get(key)
            if not isinstance(raw_values, list):
                continue
            for item in raw_values:
                marker = json.dumps(item, ensure_ascii=False, sort_keys=True)
                if marker in seen:
                    continue
                values.append(item)
                seen.add(marker)
        if values:
            merged[key] = values
    return merged


def merge_content_goals(objects: list[dict[str, Any]]) -> dict[str, int]:
    goals: dict[str, int] = {}
    for obj in objects:
        raw = obj.get("contentGoals")
        if not isinstance(raw, dict):
            continue
        for key in (
            "pages",
            "products",
            "posts",
            "navigationItems",
            "productCategories",
            "postCategories",
            "forms",
            "media",
            "siteInfoFields",
        ):
            value = raw.get(key)
            if isinstance(value, bool):
                continue
            if isinstance(value, int) and value >= 0:
                goals[key] = max(goals.get(key, 0), value)
            elif isinstance(value, str) and value.strip().isdigit():
                goals[key] = max(goals.get(key, 0), int(value.strip()))
    return goals


def build_summary(inventory: dict[str, Any], extractions: list[dict[str, Any]], args: argparse.Namespace) -> dict[str, Any]:
    extracted = [item for item in extractions if item.get("status") == "extracted"]
    failed = [item for item in extractions if item.get("status") == "failed"]
    unsupported = [item for item in extractions if item.get("status") == "unsupported"]
    text_parts = [str(item.get("text", "")) for item in extracted if item.get("text")]
    combined, combined_truncated = compact_text("\n".join(text_parts), args.max_text_chars)
    source_refs = [str(item.get("sourceRef")) for item in extracted if item.get("sourceRef")]
    source_file_fingerprints = [
        {
            "sourceRef": item.get("sourceRef"),
            "path": item.get("path"),
            "name": item.get("name"),
            "type": item.get("type"),
            "sizeBytes": item.get("sizeBytes"),
            "sha256": item.get("sha256"),
            "hashVerified": item.get("hashVerified") is True,
        }
        for item in extractions
        if isinstance(item, dict)
    ]
    fallback_title = args.site_name or "Draft AllinCMS Site"
    products = table_products(extracted)
    posts = planned_articles(extracted) or text_posts(extracted, fallback_title)
    pages = planned_pages(extracted, fallback_title, combined, source_refs)
    structured = structured_objects(extracted)
    structured_pages = merge_pages(structured)
    if structured_pages:
        pages = structured_pages
    structured_products = merge_list_fields(structured, "products")
    if structured_products:
        products = structured_products + products
    structured_posts = merge_list_fields(structured, "posts")
    if structured_posts:
        posts = structured_posts + posts
    navigation = merge_navigation(structured)
    taxonomy_plan = merge_taxonomy_plan(structured)
    content_goals = merge_content_goals(structured)
    site = merge_site(
        structured,
        {
            "siteName": fallback_title,
            "siteDescription": args.site_description or (combined[:240] if combined else "Draft site description requires source review."),
            "language": args.language,
            "industry": args.industry,
        },
    )
    return {
        "kind": "allincms_raw_extraction_summary",
        "generatedAt": now_iso(),
        "localOnly": True,
        "remoteMutationsPerformed": False,
        "inventory": args.inventory,
        "rawExtractionRefs": [str(Path(args.output_dir) / "extractions.json")],
        "wikiRefs": [str(Path(args.output_dir) / "source-wiki.json")],
        "sourceRefs": source_refs,
        "sourceFileFingerprints": source_file_fingerprints,
        "site": site,
        "pages": pages,
        "products": products,
        "posts": posts,
        "forms": merge_list_fields(structured, "forms"),
        "media": merge_list_fields(structured, "media") + [item.get("mediaCandidate") for item in extracted if isinstance(item.get("mediaCandidate"), dict)],
        "siteInfo": merge_dicts(*(obj.get("siteInfo") for obj in structured)),
        "navigation": navigation,
        "taxonomyPlan": taxonomy_plan,
        "contentGoals": content_goals,
        "mediaPolicy": merge_dicts(*(obj.get("mediaPolicy") for obj in structured)),
        "contactFormPolicy": merge_dicts(*(obj.get("contactFormPolicy") for obj in structured)),
        "openQuestions": [
            (
                "Review extracted text and create product records before final package confirmation."
                if not products
                else "Review generated product candidates against source tables before final package confirmation."
            ),
            (
                "Review extracted text and create article records before final package confirmation."
                if not posts
                else "Review generated article candidates against source text before final package confirmation."
            ),
            "Confirm public contact details, media policy, domains, and tracking before launch.",
        ],
        "extractionStats": {
            "inventoryFileCount": len(inventory.get("entries", [])) if isinstance(inventory.get("entries"), list) else 0,
            "extractedCount": len(extracted),
            "failedCount": len(failed),
            "unsupportedCount": len(unsupported),
            "hashVerifiedCount": sum(1 for item in extractions if isinstance(item, dict) and item.get("hashVerified") is True),
            "combinedTextTruncated": combined_truncated,
        },
    }


def validate_summary(summary: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    if summary.get("kind") != "allincms_raw_extraction_summary":
        issues.append("kind must be allincms_raw_extraction_summary")
    if summary.get("localOnly") is not True:
        issues.append("localOnly must be true")
    if summary.get("remoteMutationsPerformed") is not False:
        issues.append("remoteMutationsPerformed must be false")
    refs = summary.get("sourceRefs")
    if not isinstance(refs, list) or not refs:
        issues.append("sourceRefs must be non-empty")
    fingerprints = summary.get("sourceFileFingerprints")
    if not isinstance(fingerprints, list) or not fingerprints:
        issues.append("sourceFileFingerprints must be non-empty")
    elif any(isinstance(item, dict) and item.get("hashVerified") is False for item in fingerprints):
        issues.append("sourceFileFingerprints contains unverified or drifted source files")
    for text in json.dumps(summary, ensure_ascii=False).splitlines():
        for pattern in SENSITIVE_PATTERNS:
            if pattern.search(text):
                issues.append("summary contains sensitive credential/header-like text")
                return issues
    return issues


def build_extraction(args: argparse.Namespace) -> dict[str, Any]:
    output_dir = Path(args.output_dir)
    ensure_output_outside_skill(output_dir)
    inventory = load_json(Path(args.inventory), "inventory")
    inv_errors = validate_inventory(inventory)
    if inv_errors:
        raise SystemExit("ERROR: invalid inventory:\n- " + "\n- ".join(inv_errors))
    entries = inventory.get("entries")
    if not isinstance(entries, list):
        raise SystemExit("ERROR: inventory entries must be an array")
    extractions = [extract_entry(entry, args) for entry in entries if isinstance(entry, dict)]
    summary = build_summary(inventory, extractions, args)
    issues = validate_summary(summary)
    if issues:
        raise SystemExit("ERROR: invalid raw extraction summary:\n- " + "\n- ".join(issues))
    write_json(output_dir / "extractions.json", {"kind": "allincms_raw_extractions", "items": extractions})
    write_json(output_dir / "summary.json", summary)
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract source files into raw extraction artifacts for AllinCMS packaging.")
    parser.add_argument("--inventory", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--site-name", default="")
    parser.add_argument("--site-description", default="")
    parser.add_argument("--language", default="en")
    parser.add_argument("--industry", default="unspecified")
    parser.add_argument("--max-text-chars", type=int, default=MAX_TEXT_CHARS_DEFAULT)
    parser.add_argument("--max-table-rows", type=int, default=MAX_TABLE_ROWS_DEFAULT)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    summary = build_extraction(args)
    print(f"Wrote raw extraction artifacts under: {args.output_dir}")
    stats = summary["extractionStats"]
    print(f"extracted={stats['extractedCount']} failed={stats['failedCount']} unsupported={stats['unsupportedCount']}")
    if args.json:
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
